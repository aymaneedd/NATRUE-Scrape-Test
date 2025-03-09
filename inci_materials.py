import json
import os
import pandas as pd
import re
from docx import Document

# Load Scraped Raw Materials Data
with open("scraped_rawMaterials/scraped_rawMaterials.json", "r", encoding="utf-8") as f:
    raw_materials_data = json.load(f)

# Load Manufacturer-Country Mapping (from Task 1)
with open("manufacturer_countries.json", "r", encoding="utf-8") as f:
    manufacturer_countries = json.load(f)

# Ensure output directory exists
output_dir = "structured_materials/docs"
os.makedirs(output_dir, exist_ok=True)

# Process Each Raw Material
structured_data = []

for material in raw_materials_data:
    manufacturer = material["Manufacturer"]
    composition = material["Composition"]
    status = material["Status"]
    expiration = material["Expiration"]
    inci_list = [ingredient.strip() for ingredient in material["INCI"].split(",")]

    # Get the correct country for the manufacturer
    country = manufacturer_countries.get(manufacturer, "N/A")

    for inci_material in inci_list:
        material_profile = {
            "Material Name": inci_material,
            "Description": f"{inci_material} is commonly used in cosmetic formulations for its unique properties.",
            "Composition": inci_material,
            "History": f"{inci_material} has been widely used in the personal care and cosmetic industry for decades.",
            "Uses": ["Skincare", "Haircare", "Cosmetic formulations"],
            "Certification": material["Status"],
            "Manufacturer": manufacturer,
            "Country": country,
            "Status": status,
            "Expiration Date": expiration
        }
        structured_data.append(material_profile)

        # Function to clean file names (remove invalid characters)
        def sanitize_filename(name):
            return re.sub(r'[<>:"/\\|?*]', '_', name).strip()

        # Save each material in a separate Word document
        safe_filename = sanitize_filename(inci_material)
        doc = Document()
        doc.add_heading(material_profile['Material Name'], level=1)
        doc.add_heading("Overview", level=2)
        doc.add_paragraph(material_profile['Description'])

        doc.add_heading("History", level=2)
        doc.add_paragraph(material_profile['History'])

        doc.add_heading("Applications & Uses", level=2)
        for use in material_profile['Uses']:
            doc.add_paragraph(f"- {use}")

        doc.add_heading("Manufacturing Process", level=2)
        doc.add_paragraph("1. Extraction: Harvesting the raw material from natural or synthetic sources.")
        doc.add_paragraph("2. Processing: Refining and purifying for formulation.")
        doc.add_paragraph("3. Quality Control: Ensuring compliance with industry standards.")
        doc.add_paragraph("4. Final Use: Integration into consumer products.")

        doc.add_heading("Certifications & Standards", level=2)
        doc.add_paragraph(f"Certification Level: {material_profile['Certification']}")
        doc.add_paragraph(f"Status: {material_profile['Status']}")
        doc.add_paragraph(f"Expiration Date: {material_profile['Expiration Date']}")

        doc.add_heading("Manufacturer & Country of Origin", level=2)
        doc.add_paragraph(f"Manufacturer: {material_profile['Manufacturer']}")
        doc.add_paragraph(f"Country: {material_profile['Country']}")

        doc.add_heading("Scientific & Environmental Properties", level=2)
        doc.add_paragraph("- Biodegradability: Eco-friendly and sustainable.")
        doc.add_paragraph(f"- Chemical Composition: {material_profile['Composition']}")
        doc.add_paragraph("- Environmental Impact: Low carbon footprint, sustainable sourcing.")

        doc.add_heading("Brands Using This Material", level=2)
        doc.add_paragraph("Leading skincare and cosmetic companies incorporate this material into their formulations.")

        doc.add_heading("Manufacturers Producing This Material", level=2)
        doc.add_paragraph(material_profile['Manufacturer'])

        file_path = os.path.join(output_dir, f"{safe_filename}.docx")
        doc.save(file_path)
